import pandas as pd
import tkinter as tk
import numpy as np
import docx as word
from datetime import date

#specifying date
date=date.today()
#main window creation
tab=tk.Tk()
tab.title("SALES MANAGEMENT SYSTEM")
tab.geometry("1000x600")
tk.Label(tab,text="========WELCOME TO SALES MANAGEMENT SYSTEM!========",font=("Arial rounded MT bold",13)).grid(row=0,column=1)
tk.Label(tab,text="WOULD YOUR LIKE TO START RECORDING TRANSACTIONS?",font=("Arial rounded MT bold",12)).grid(row=1,column=0)
tk.Label(tab,text="IF YES CLICK THIS BUTTON:",font=("Arial rounded MT bold",10)).grid(row=2,column=0)
tk.Label(tab,text="IF NO CLICK THIS BUTTON:",font=("Arial rounded MT bold",10)).grid(row=3,column=0)
tk.Label(tab,text="SHOW DATA OF PARTICULAR DATE:",font=("Arial rounded MT bold",10)).grid(row=4,column=0)
tk.Label(tab,text="MODIFY/DELETE A RECORD:",font=("Arial rounded MT bold",10)).grid(row=5,column=0)
tk.Label(tab,text="SHOW ENTIRE DATA:",font=("Arial rounded MT bold",10)).grid(row=6,column=0)
tk.Label(tab,text="SHOW INVENTORY:",font=("Arial rounded MT bold",10)).grid(row=7,column=0)
#Assinging string which is to be converted for further processing.
name=tk.StringVar()
phone=tk.StringVar()
itemcode=tk.StringVar()
qtn=tk.StringVar()
billno=tk.StringVar()
discount=tk.IntVar()
def YES():
    tk.Label(tab,text="ENTER NAME OF THE CUSTOMER:",font=("Arial rounded MT bold",10)).grid(row=5,column=0)
    tk.Entry(tab,textvariable=name).grid(row=5,column=1)
    tk.Label(tab,text="ENTER PHONE NUMBER OF THE CUSTOMER:",font=("Arial rounded MT bold",10)).grid(row=6,column=0)
    tk.Entry(tab,textvariable=phone).grid(row=6,column=1)
    tk.Label(tab,text="ENTER THE ITEMCODE:",font=("Arial rounded MT bold",10)).grid(row=7,column=0)
    tk.Entry(tab,textvariable=itemcode).grid(row=7,column=1)
    tk.Label(tab,text="ENTER THE QUANTITY OF THE PRODUCTS PURCHASED:",font=("Arial rounded MT bold",10)).grid(row=8,column=0)
    tk.Entry(tab,textvariable=qtn).grid(row=8,column=1)
    tk.Label(tab,text="ENTER THE BILL NO:",font=("Arial rounded MT bold",10)).grid(row=9,column=0)
    tk.Entry(tab,textvariable=billno).grid(row=9,column=1)
    tk.Label(tab,text="GIVE DISCOUNT?(0 FOR NO)",font=("Arial rounded MT bold",10)).grid(row=10,column=0)
    tk.Entry(tab,textvariable=discount).grid(row=10,column=1)
    #function below converts the previously assinged strings to stated data types for further processing(UPDATE).
    def record():
        name1=name.get()
        phone1=int(phone.get())
        itemcode1=int(itemcode.get())
        qtn1=int(qtn.get())
        billno1=int(billno.get())
        discount1=int(discount.get())
        stock=pd.read_csv("D:\DATA\dataSTOCK")
        temp1=list(stock["itemcode"])
        for i in range(9):
            if temp1[i]==itemcode1:
                temp2=stock.loc[[i],["price"]]
                temp4=list(stock["itemname"])
        price=int(temp2["price"])
        name2=str(temp4[itemcode1])
        Amount=price*qtn1
        netamount=Amount-(Amount/100*discount1)
        netamount=int(netamount)
        #storing/updating "sales" csv i.e., adding a new record.
        sales=pd.read_csv("D:\DATA\dataSALES",usecols=["billno","nameofcustomer","phone","date","name","item","price","qtn","totalamount","discount","netamount"])
        print(sales)
        sales.loc[billno1,:]=[billno1,name1,phone1,date,name2,itemcode1,price,qtn1,Amount,discount1,netamount]
        print("AFTER UPDATING THE CSV FILE")
        print("\n")
        print(sales)
        sales.to_csv("D:\DATA\dataSALES")
        #updating "stock" csv i.e., reducing from the quantity.
        for i in range(9):
            if temp1[i]==itemcode1:
                stock.loc[[i],["qtn"]]-=qtn1
                stock.to_csv("D:\DATA\dataStock")
        #Generated bill
        print("|-------------XYZ HELLO ENTERPRISE---------------") 
        print("|======= Shopno 101 ,whiterun, skyrim, tamrial =======")
        print("|Billno:",billno1)
        print("|NAME OF THE CUSTOMER:",name1,"            Date:",date)
        print("|Phone no:",phone1)
        print("|itemcode|Name|price   |Quantity|Amount    |Discount|netamount |")
        print("|",itemcode1,"    |",name2,"|",price,"|",qtn1,"    |",Amount,"|",discount1,"%   |",netamount,"|")
        print("|Total amount",netamount)
        #Printable bill
        doc=word.Document("D:\DATA\print_bill.docx")
        doc.add_paragraph("|   -------------XYZ HELLO ENTERPRISE--------------")
        doc.add_paragraph("|======= Shopno 101 ,whiterun, skyrim, tamrial =======")
        doc.add_paragraph("|BILLno:")
        doc.add_paragraph(str(billno1))
        doc.add_paragraph("|NAME OF THE CUSTOMER:")
        doc.add_paragraph(str(name1))
        doc.add_paragraph("|Phone no:")
        doc.add_paragraph(str(phone1))
        doc.add_paragraph("|DATE:")
        doc.add_paragraph(str(date))
        doc.add_paragraph("|=======================================================")
        doc.add_paragraph("|Itemcode:")
        doc.add_paragraph(str(itemcode1))
        doc.add_paragraph("|Name:")
        doc.add_paragraph(str(name2))
        doc.add_paragraph("|Price:")
        doc.add_paragraph(str(price))
        doc.add_paragraph("|Quantity:")
        doc.add_paragraph(str(qtn1))
        doc.add_paragraph("|Amount:")
        doc.add_paragraph(str(Amount))
        doc.add_paragraph("|Discount:")
        doc.add_paragraph(str(discount1))
        doc.add_paragraph("|Netamount:")
        doc.add_paragraph(str(netamount))
        doc.add_paragraph("|Total Amount:")
        doc.add_paragraph(str(netamount))
        doc.save("D:\DATA\print_bill.docx")
    tk.Button(tab,text="Record",command=record).grid(row=12,column=1)
tk.Button(tab,text="YES",command=YES).grid(row=2,column=1)
def Exit():
    quit()
tk.Button(tab,text="EXIT",command=Exit).grid(row=3,column=1)
# output no 2 and 3
def COMMENCE():
    sales=pd.read_csv("D:\DATA\dataSALES",usecols=["billno","nameofcustomer","phone","date","name","item","price","qtn","totalamount","discount","netamount"])
    def ENTER_DATE():
        date=tk.StringVar()
        tk.Label(tab,text="ENTER THE DATE OF CHOICE(YYYY-MM-DD):",font=("Arial rounded MT bold",10)).grid(row=5,column=0)
        tk.Entry(tab,textvariable=date).grid(row=5,column=1)
        def SHOW_dis():
            date1=str(date.get())
            print("DATE:",date1)
            temp6=sales.loc[:,["billno","nameofcustomer","item","discount"]]
            temp7=sales["date"]
            for i in range(len(temp6)):
                if str(temp7[i])==date1:
                    print(temp6.loc[[i],["billno","nameofcustomer","item","discount"]])
                else:
                    print("")
        tk.Button(tab,text="SHOW ALL DISCOUNTS",command=SHOW_dis).grid(row=8,column=1)            
        def SHOW_pur():
            date1=str(date.get())
            print("DATE:",date1)
            temp6=sales.loc[:,["billno","nameofcustomer","item","netamount"]]
            temp7=sales["date"]
            for i in range(len(temp6)):
                if str(temp7[i])==date1:
                    print(temp6.loc[[i],["billno","nameofcustomer","item","netamount"]])
                else:
                    print("")
        tk.Button(tab,text="SHOW ALL SALES",command=SHOW_pur).grid(row=7,column=1)            
    tk.Button(tab,text="ENTER DATE",command=ENTER_DATE).grid(row=6,column=1)            
tk.Button(tab,text="COMMENCE",command=COMMENCE).grid(row=4,column=1)
#To modify/delete a single record.
def mod_del():
    sales=pd.read_csv("D:\DATA\dataSALES",usecols=["billno","nameofcustomer","phone","date","name","item","price","qtn","totalamount","discount","netamount"])
    def ENTER_BILL():
        bill_temp=tk.StringVar()
        tk.Label(tab,text="ENTER THE VALID BILL NUMBER OF CHOICE:",font=("Arial rounded MT bold",10)).grid(row=6,column=0)
        tk.Entry(tab,textvariable=bill_temp).grid(row=6,column=1)      
        def mod():
            namemod=tk.StringVar()
            billmod=tk.StringVar()
            phonemod=tk.StringVar()
            tk.Label(tab,text="ENTER NEW NAME(MUST NOT BE EMPTY):",font=("Arial rounded MT bold",10)).grid(row=7,column=0)
            tk.Entry(tab,textvariable=namemod).grid(row=7,column=1)
            tk.Label(tab,text="ENTER NEW BILL(MUST NOT BE EMPTY):",font=("Arial rounded MT bold",10)).grid(row=8,column=0)
            tk.Entry(tab,textvariable=billmod).grid(row=8,column=1)
            tk.Label(tab,text="ENTER NEW PHONE(MUST NOT BE EMPTY):",font=("Arial rounded MT bold",10)).grid(row=9,column=0)
            tk.Entry(tab,textvariable=phonemod).grid(row=9,column=1)
            sales=pd.read_csv("D:\DATA\dataSALES",usecols=["billno","nameofcustomer","phone","date","name","item","price","qtn","totalamount","discount","netamount"])
            def done_mod():
                modname=str(namemod.get())
                modbill=int(billmod.get())
                modphone=str(phonemod.get())
                temp8=int(bill_temp.get())
                print("DONE")
                sales.loc[[temp8-1],["billno","nameofcustomer","phone"]]=[modbill,modname,modphone]
                sales.to_csv("D:\DATA\dataSALES")
            tk.Button(tab,text="DONE",command=done_mod).grid(row=10,column=1)    
        tk.Button(tab,text="MODIFY",command=mod).grid(row=8,column=1)
        def DEL():
            temp8=int(bill_temp.get())
            print("DONE") 
            sales.drop([temp8],axis=0,inplace=True)
            sales.to_csv("D:\DATA\dataSALES")
        tk.Button(tab,text="DELETE",command=DEL).grid(row=9,column=1)    
    tk.Button(tab,text="ENTER BILL",command=ENTER_BILL).grid(row=7,column=1)        
tk.Button(tab,text="COMMENCE",command=mod_del).grid(row=5,column=1)
#to display stored data
def SHOW_data():
    sales=pd.read_csv("D:\DATA\dataSALES",usecols=["billno","nameofcustomer","phone","date","name","item","price","qtn","totalamount","discount","netamount"])
    print(sales)
tk.Button(tab,text="DISPLAY",command=SHOW_data).grid(row=6,column=1)
def SHOW_inventory():
    stock=pd.read_csv("D:\DATA\dataSTOCK",usecols=["itemcode","itemname","price","qtn"])
    print(stock)
tk.Button(tab,text="DISPLAY",command=SHOW_inventory).grid(row=7,column=1)    
tab.mainloop()

